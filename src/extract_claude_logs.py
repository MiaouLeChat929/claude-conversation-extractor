#!/usr/bin/env python3
"""
Extract clean conversation logs from Claude Code's internal JSONL files

This tool parses the undocumented JSONL format used by Claude Code to store
conversations locally in ~/.claude/projects/ and exports them as clean,
readable markdown files.
"""

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Optional dependencies for additional export formats
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.colors import HexColor

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import shared constants
try:
    from .constants import (
        INDENT_NUMBER,
        MAJOR_SEPARATOR_WIDTH,
        MINOR_SEPARATOR_WIDTH,
        SESSION_ID_MAX_LENGTH,
        LINES_SHOWN_MESSAGE,
        LINES_PER_PAGE_MESSAGE,
        MAX_LINES_PER_MESSAGE_DISPLAY,
        MAX_LINE_LENGTH_DISPLAY,
        MIN_PREVIEW_TEXT_LENGTH,
        PREVIEW_TEXT_TRUNCATE_LENGTH,
        PREVIEW_ERROR_TRUNCATE_LENGTH,
        LIST_SEPARATOR_WIDTH,
        SEARCH_MAX_RESULTS_DEFAULT,
    )
except ImportError:
    from constants import (
        INDENT_NUMBER,
        MAJOR_SEPARATOR_WIDTH,
        MINOR_SEPARATOR_WIDTH,
        SESSION_ID_MAX_LENGTH,
        LINES_SHOWN_MESSAGE,
        LINES_PER_PAGE_MESSAGE,
        MAX_LINES_PER_MESSAGE_DISPLAY,
        MAX_LINE_LENGTH_DISPLAY,
        MIN_PREVIEW_TEXT_LENGTH,
        PREVIEW_TEXT_TRUNCATE_LENGTH,
        PREVIEW_ERROR_TRUNCATE_LENGTH,
        LIST_SEPARATOR_WIDTH,
        SEARCH_MAX_RESULTS_DEFAULT,
    )


class ClaudeConversationExtractor:
    """Extract and convert Claude Code conversations from JSONL to markdown."""

    def __init__(self, output_dir: Optional[Path] = None):
        """Initialize the extractor with Claude's directory and output location."""
        self.claude_dir = Path.home() / ".claude" / "projects"

        if output_dir:
            self.output_dir = Path(output_dir)
            self.output_dir.mkdir(parents=True, exist_ok=True)
        else:
            # Try multiple possible output directories
            possible_dirs = [
                Path.home() / "Desktop" / "Claude logs",
                Path.home() / "Documents" / "Claude logs",
                Path.home() / "Claude logs",
                Path.cwd() / "claude-logs",
            ]

            # Use the first directory we can create
            for dir_path in possible_dirs:
                try:
                    dir_path.mkdir(parents=True, exist_ok=True)
                    # Test if we can write to it
                    test_file = dir_path / ".test"
                    test_file.touch()
                    test_file.unlink()
                    self.output_dir = dir_path
                    break
                except Exception:
                    continue
            else:
                # Fallback to current directory
                self.output_dir = Path.cwd() / "claude-logs"
                self.output_dir.mkdir(exist_ok=True)

        print(f"📁 Saving logs to: {self.output_dir}")

    def find_sessions(self, project_path: Optional[str] = None) -> List[Path]:
        """Find all JSONL session files, sorted by most recent first."""
        if project_path:
            search_dir = self.claude_dir / project_path
        else:
            search_dir = self.claude_dir

        sessions = []
        if search_dir.exists():
            for jsonl_file in search_dir.rglob("*.jsonl"):
                sessions.append(jsonl_file)
        return sorted(sessions, key=lambda x: x.stat().st_mtime, reverse=True)

    def extract_conversation(
        self, jsonl_path: Path, detailed: bool = False
    ) -> List[Dict[str, str]]:
        """Extract conversation messages from a JSONL file.

        Args:
            jsonl_path: Path to the JSONL file
            detailed: If True, include tool use, MCP responses, and system messages
        """
        conversation = []

        try:
            with open(jsonl_path, "r", encoding="utf-8") as f:
                for line in f:
                    try:
                        entry = json.loads(line.strip())

                        # Extract user messages
                        if entry.get("type") == "user" and "message" in entry:
                            msg = entry["message"]
                            if isinstance(msg, dict) and msg.get("role") == "user":
                                content = msg.get("content", "")
                                text = self._extract_text_content(content)

                                if text and text.strip():
                                    conversation.append(
                                        {
                                            "role": "user",
                                            "content": text,
                                            "timestamp": entry.get("timestamp", ""),
                                        }
                                    )

                        # Extract assistant messages
                        elif entry.get("type") == "assistant" and "message" in entry:
                            msg = entry["message"]
                            if isinstance(msg, dict) and msg.get("role") == "assistant":
                                content = msg.get("content", [])
                                text = self._extract_text_content(content, detailed=detailed)

                                if text and text.strip():
                                    conversation.append(
                                        {
                                            "role": "assistant",
                                            "content": text,
                                            "timestamp": entry.get("timestamp", ""),
                                        }
                                    )

                        # Include tool use and system messages if detailed mode
                        elif detailed:
                            # Extract tool use events
                            if entry.get("type") == "tool_use":
                                tool_data = entry.get("tool", {})
                                tool_name = tool_data.get("name", "unknown")
                                tool_input = tool_data.get("input", {})
                                conversation.append(
                                    {
                                        "role": "tool_use",
                                        "content": f"🔧 Tool: {tool_name}\nInput: {json.dumps(tool_input, indent=INDENT_NUMBER, ensure_ascii=False)}",
                                        "timestamp": entry.get("timestamp", ""),
                                    }
                                )

                            # Extract tool results
                            elif entry.get("type") == "tool_result":
                                result = entry.get("result", {})
                                output = result.get("output", "") or result.get("error", "")
                                conversation.append(
                                    {
                                        "role": "tool_result",
                                        "content": f"📤 Result:\n{output}",
                                        "timestamp": entry.get("timestamp", ""),
                                    }
                                )

                            # Extract system messages
                            elif entry.get("type") == "system" and "message" in entry:
                                msg = entry.get("message", "")
                                if msg:
                                    conversation.append(
                                        {
                                            "role": "system",
                                            "content": f"ℹ️ System: {msg}",
                                            "timestamp": entry.get("timestamp", ""),
                                        }
                                    )

                    except json.JSONDecodeError:
                        continue
                    except Exception:
                        # Silently skip problematic entries
                        continue

        except Exception as e:
            print(f"❌ Error reading file {jsonl_path}: {e}")

        return conversation

    def _extract_text_content(self, content, detailed: bool = False) -> str:
        """Extract text from various content formats Claude uses.

        Args:
            content: The content to extract from
            detailed: If True, include tool use blocks and other metadata
        """
        if isinstance(content, str):
            return content
        elif isinstance(content, list):
            # Extract text from content array
            text_parts = []
            for item in content:
                if isinstance(item, dict):
                    if item.get("type") == "text":
                        text_parts.append(item.get("text", ""))
                    elif detailed and item.get("type") == "tool_use":
                        # Include tool use details in detailed mode
                        tool_name = item.get("name", "unknown")
                        tool_input = item.get("input", {})
                        text_parts.append(f"\n🔧 Using tool: {tool_name}")
                        text_parts.append(
                            f"Input: {json.dumps(tool_input, indent=INDENT_NUMBER, ensure_ascii=False)}\n"
                        )
            return "\n".join(text_parts)
        else:
            return str(content)

    def _parse_timestamp(self, conversation: List[Dict[str, str]]) -> Tuple[str, str]:
        """Parse timestamp from conversation to get date and time strings.

        Args:
            conversation: List of message dictionaries

        Returns:
            Tuple of (date_str, time_str)
        """
        if not conversation:
            return datetime.now().strftime("%Y-%m-%d"), ""

        first_timestamp = conversation[0].get("timestamp", "")
        if first_timestamp:
            try:
                dt = datetime.fromisoformat(first_timestamp.replace("Z", "+00:00"))
                return dt.strftime("%Y-%m-%d"), dt.strftime("%H:%M:%S")
            except Exception:
                pass
        return datetime.now().strftime("%Y-%m-%d"), ""

    def _clean_project_name(self, project_name: str) -> str:
        """Clean project name for use in filenames.

        Args:
            project_name: Raw project name

        Returns:
            Sanitized project name safe for filenames
        """
        if not project_name:
            return ""
        return project_name.replace("/", "-").replace("\\", "-").replace(" ", "-")

    def _generate_filename(
        self, project_name: str, date_str: str, session_id: str, extension: str
    ) -> str:
        """Generate a filename for exported conversation.

        Args:
            project_name: Cleaned project name (can be empty)
            date_str: Date string in YYYY-MM-DD format
            session_id: Session identifier
            extension: File extension (e.g., 'md', 'json', 'html')

        Returns:
            Generated filename
        """
        clean_project = self._clean_project_name(project_name)
        if clean_project:
            return f"claude-conversation-{clean_project}-{date_str}-{session_id[:SESSION_ID_MAX_LENGTH]}.{extension}"
        return f"claude-conversation-{date_str}-{session_id[:SESSION_ID_MAX_LENGTH]}.{extension}"

    def display_conversation(self, jsonl_path: Path, detailed: bool = False) -> None:
        """Display a conversation in the terminal with pagination.

        Args:
            jsonl_path: Path to the JSONL file
            detailed: If True, include tool use and system messages
        """
        try:
            # Extract conversation
            messages = self.extract_conversation(jsonl_path, detailed=detailed)

            if not messages:
                print("❌ No messages found in conversation")
                return

            # Get session info
            session_id = jsonl_path.stem

            # Clear screen and show header
            print("\033[2J\033[H", end="")  # Clear screen
            print("=" * MAJOR_SEPARATOR_WIDTH)
            print(f"📄 Viewing: {jsonl_path.parent.name}")
            print(f"Session: {session_id[:SESSION_ID_MAX_LENGTH]}...")

            # Get timestamp from first message
            first_timestamp = messages[0].get("timestamp", "")
            if first_timestamp:
                try:
                    dt = datetime.fromisoformat(first_timestamp.replace("Z", "+00:00"))
                    print(f"Date: {dt.strftime('%Y-%m-%d %H:%M:%S')}")
                except Exception:
                    pass

            print("=" * MAJOR_SEPARATOR_WIDTH)
            print("↑↓ to scroll • Q to quit • Enter to continue\n")

            # Display messages with pagination
            lines_shown = LINES_SHOWN_MESSAGE  # Header lines
            lines_per_page = LINES_PER_PAGE_MESSAGE

            for i, msg in enumerate(messages):
                role = msg["role"]
                content = msg["content"]

                # Format role display
                if role == "user" or role == "human":
                    print(f"\n{'─' * MINOR_SEPARATOR_WIDTH}")
                    print(f"👤 HUMAN:")
                    print(f"{'─' * MINOR_SEPARATOR_WIDTH}")
                elif role == "assistant":
                    print(f"\n{'─' * MINOR_SEPARATOR_WIDTH}")
                    print(f"🤖 CLAUDE:")
                    print(f"{'─' * MINOR_SEPARATOR_WIDTH}")
                elif role == "tool_use":
                    print(f"\n🔧 TOOL USE:")
                elif role == "tool_result":
                    print(f"\n📤 TOOL RESULT:")
                elif role == "system":
                    print(f"\nℹ️ SYSTEM:")
                else:
                    print(f"\n{role.upper()}:")

                # Display content (limit very long messages)
                lines = content.split("\n")
                max_lines_per_msg = MAX_LINES_PER_MESSAGE_DISPLAY

                for line_idx, line in enumerate(lines[:max_lines_per_msg]):
                    # Wrap very long lines
                    if len(line) > MAX_LINE_LENGTH_DISPLAY:
                        line = line[: (MAX_LINE_LENGTH_DISPLAY - 3)] + "..."
                    print(line)
                    lines_shown += 1

                    # Check if we need to paginate
                    if lines_shown >= lines_per_page:
                        response = input("\n[Enter] Continue • [Q] Quit: ").strip().upper()
                        if response == "Q":
                            print("\n👋 Stopped viewing")
                            return
                        # Clear screen for next page
                        print("\033[2J\033[H", end="")
                        lines_shown = 0

                if len(lines) > max_lines_per_msg:
                    print(f"... [{len(lines) - max_lines_per_msg} more lines truncated]")
                    lines_shown += 1

            print("\n" + "=" * MAJOR_SEPARATOR_WIDTH)
            print("📄 End of conversation")
            print("=" * MAJOR_SEPARATOR_WIDTH)
            input("\nPress Enter to continue...")

        except Exception as e:
            print(f"❌ Error displaying conversation: {e}")
            input("\nPress Enter to continue...")

    def save_as_markdown(
        self, conversation: List[Dict[str, str]], session_id: str, project_name: str = ""
    ) -> Optional[Path]:
        """Save conversation as clean markdown file."""
        if not conversation:
            return None

        date_str, time_str = self._parse_timestamp(conversation)
        filename = self._generate_filename(project_name, date_str, session_id, "md")
        output_path = self.output_dir / filename

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("# Claude Conversation Log\n\n")
            f.write(f"Session ID: {session_id}\n")
            f.write(f"Date: {date_str}")
            if time_str:
                f.write(f" {time_str}")
            f.write("\n\n---\n\n")

            for msg in conversation:
                role = msg["role"]
                content = msg["content"]

                if role == "user":
                    f.write("## 👤 User\n\n")
                    f.write(f"{content}\n\n")
                elif role == "assistant":
                    f.write("## 🤖 Claude\n\n")
                    f.write(f"{content}\n\n")
                elif role == "tool_use":
                    f.write("### 🔧 Tool Use\n\n")
                    f.write(f"{content}\n\n")
                elif role == "tool_result":
                    f.write("### 📤 Tool Result\n\n")
                    f.write(f"{content}\n\n")
                elif role == "system":
                    f.write("### ℹ️ System\n\n")
                    f.write(f"{content}\n\n")
                else:
                    f.write(f"## {role}\n\n")
                    f.write(f"{content}\n\n")
                f.write("---\n\n")

        return output_path

    def save_as_json(
        self, conversation: List[Dict[str, str]], session_id: str, project_name: str = ""
    ) -> Optional[Path]:
        """Save conversation as JSON file."""
        if not conversation:
            return None

        date_str, _ = self._parse_timestamp(conversation)
        filename = self._generate_filename(project_name, date_str, session_id, "json")
        output_path = self.output_dir / filename

        # Create JSON structure
        output = {
            "session_id": session_id,
            "date": date_str,
            "message_count": len(conversation),
            "messages": conversation,
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=INDENT_NUMBER, ensure_ascii=False)

        return output_path

    def save_as_html(
        self, conversation: List[Dict[str, str]], session_id: str, project_name: str = ""
    ) -> Optional[Path]:
        """Save conversation as HTML file with syntax highlighting."""
        if not conversation:
            return None

        date_str, time_str = self._parse_timestamp(conversation)
        filename = self._generate_filename(project_name, date_str, session_id, "html")
        output_path = self.output_dir / filename

        # HTML template with modern styling
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Claude Conversation - {session_id[:SESSION_ID_MAX_LENGTH]}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .header {{
            background: white;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 20px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            margin: 0 0 10px 0;
        }}
        .metadata {{
            color: #666;
            font-size: 0.9em;
        }}
        .message {{
            background: white;
            padding: 15px 20px;
            margin-bottom: 15px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .user {{
            border-left: 4px solid #3498db;
        }}
        .assistant {{
            border-left: 4px solid #2ecc71;
        }}
        .tool_use {{
            border-left: 4px solid #f39c12;
            background: #fffbf0;
        }}
        .tool_result {{
            border-left: 4px solid #e74c3c;
            background: #fff5f5;
        }}
        .system {{
            border-left: 4px solid #95a5a6;
            background: #f8f9fa;
        }}
        .role {{
            font-weight: bold;
            margin-bottom: 10px;
            display: flex;
            align-items: center;
        }}
        .content {{
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        pre {{
            background: #f4f4f4;
            padding: 10px;
            border-radius: 4px;
            overflow-x: auto;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Claude Conversation Log</h1>
        <div class="metadata">
            <p>Session ID: {session_id}</p>
            <p>Date: {date_str} {time_str}</p>
            <p>Messages: {len(conversation)}</p>
        </div>
    </div>
"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

            for msg in conversation:
                role = msg["role"]
                content = msg["content"]

                # Escape HTML
                content = content.replace("&", "&amp;")
                content = content.replace("<", "&lt;")
                content = content.replace(">", "&gt;")

                role_display = {
                    "user": "👤 User",
                    "assistant": "🤖 Claude",
                    "tool_use": "🔧 Tool Use",
                    "tool_result": "📤 Tool Result",
                    "system": "ℹ️ System",
                }.get(role, role)

                f.write(f'    <div class="message {role}">\n')
                f.write(f'        <div class="role">{role_display}</div>\n')
                f.write(f'        <div class="content">{content}</div>\n')
                f.write(f"    </div>\n")

            f.write("\n</body>\n</html>")

        return output_path

    def save_as_pdf(
        self, conversation: List[Dict[str, str]], session_id: str, project_name: str = ""
    ) -> Optional[Path]:
        """Save conversation as PDF file.

        Requires: pip install claude-conversation-extractor[pdf]
        """
        if not PDF_AVAILABLE:
            print("❌ PDF export requires reportlab. Install with:")
            print("   pip install claude-conversation-extractor[pdf]")
            return None

        if not conversation:
            return None

        date_str, time_str = self._parse_timestamp(conversation)
        filename = self._generate_filename(project_name, date_str, session_id, "pdf")
        output_path = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=12,
            textColor=HexColor("#2c3e50"),
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=10,
            textColor=HexColor("#666666"),
            spaceAfter=6,
        )
        user_style = ParagraphStyle(
            "User",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=HexColor("#3498db"),
            spaceBefore=12,
            spaceAfter=6,
        )
        assistant_style = ParagraphStyle(
            "Assistant",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=HexColor("#2ecc71"),
            spaceBefore=12,
            spaceAfter=6,
        )
        content_style = ParagraphStyle(
            "Content",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=12,
        )

        # Build story (content)
        story = []

        # Title
        story.append(Paragraph("Claude Conversation Log", title_style))
        story.append(Paragraph(f"Session ID: {session_id}", meta_style))
        story.append(Paragraph(f"Date: {date_str} {time_str}", meta_style))
        story.append(Paragraph(f"Messages: {len(conversation)}", meta_style))
        story.append(Spacer(1, 12))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cccccc")))
        story.append(Spacer(1, 12))

        # Messages
        for msg in conversation:
            role = msg["role"]
            content = msg["content"]

            # Escape special characters for PDF
            content = content.replace("&", "&amp;")
            content = content.replace("<", "&lt;")
            content = content.replace(">", "&gt;")
            # Replace newlines with <br/> for PDF
            content = content.replace("\n", "<br/>")

            if role == "user":
                story.append(Paragraph("👤 User", user_style))
            elif role == "assistant":
                story.append(Paragraph("🤖 Claude", assistant_style))
            elif role == "tool_use":
                story.append(Paragraph("🔧 Tool Use", user_style))
            elif role == "tool_result":
                story.append(Paragraph("📤 Tool Result", user_style))
            elif role == "system":
                story.append(Paragraph("ℹ️ System", meta_style))
            else:
                story.append(Paragraph(role, user_style))

            story.append(Paragraph(content, content_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#eeeeee")))

        doc.build(story)
        return output_path

    def save_as_docx(
        self, conversation: List[Dict[str, str]], session_id: str, project_name: str = ""
    ) -> Optional[Path]:
        """Save conversation as DOCX file.

        Requires: pip install claude-conversation-extractor[docx]
        """
        if not DOCX_AVAILABLE:
            print("❌ DOCX export requires python-docx. Install with:")
            print("   pip install claude-conversation-extractor[docx]")
            return None

        if not conversation:
            return None

        date_str, time_str = self._parse_timestamp(conversation)
        filename = self._generate_filename(project_name, date_str, session_id, "docx")
        output_path = self.output_dir / filename

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("Claude Conversation Log", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Session ID: {session_id}\n").bold = True
        meta.add_run(f"Date: {date_str} {time_str}\n")
        meta.add_run(f"Messages: {len(conversation)}")

        doc.add_paragraph("_" * 50)

        # Messages
        for msg in conversation:
            role = msg["role"]
            content = msg["content"]

            # Add role heading with appropriate color
            role_para = doc.add_paragraph()
            role_run = role_para.add_run()

            if role == "user":
                role_run.text = "👤 User"
                role_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
            elif role == "assistant":
                role_run.text = "🤖 Claude"
                role_run.font.color.rgb = RGBColor(46, 204, 113)  # Green
            elif role == "tool_use":
                role_run.text = "🔧 Tool Use"
                role_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
            elif role == "tool_result":
                role_run.text = "📤 Tool Result"
                role_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
            elif role == "system":
                role_run.text = "ℹ️ System"
                role_run.font.color.rgb = RGBColor(149, 165, 166)  # Gray
            else:
                role_run.text = role

            role_run.bold = True
            role_run.font.size = Pt(12)

            # Add content
            content_para = doc.add_paragraph(content)
            content_para.paragraph_format.space_after = Pt(12)

            # Add separator
            doc.add_paragraph("─" * 40)

        doc.save(str(output_path))
        return output_path

    def save_conversation(
        self,
        conversation: List[Dict[str, str]],
        session_id: str,
        format: str = "markdown",
        project_name: str = "",
    ) -> Optional[Path]:
        """Save conversation in the specified format.

        Args:
            conversation: The conversation data
            session_id: Session identifier
            format: Output format ('markdown', 'json', 'html', 'pdf', 'docx')
            project_name: Project name to include in filename and content
        """
        if format == "markdown":
            return self.save_as_markdown(conversation, session_id, project_name)
        elif format == "json":
            return self.save_as_json(conversation, session_id, project_name)
        elif format == "html":
            return self.save_as_html(conversation, session_id, project_name)
        elif format == "pdf":
            return self.save_as_pdf(conversation, session_id, project_name)
        elif format == "docx":
            return self.save_as_docx(conversation, session_id, project_name)
        else:
            print(f"❌ Unsupported format: {format}")
            return None

    def get_conversation_preview(self, session_path: Path) -> Tuple[str, int]:
        """Get a preview of the conversation's first real user message and message count."""
        try:
            first_user_msg = ""
            msg_count = 0

            with open(session_path, "r", encoding="utf-8") as f:
                for line in f:
                    msg_count += 1
                    if not first_user_msg:
                        try:
                            data = json.loads(line)
                            # Check for user message
                            if data.get("type") == "user" and "message" in data:
                                msg = data["message"]
                                if msg.get("role") == "user":
                                    content = msg.get("content", "")

                                    # Handle list content (common format in Claude JSONL)
                                    if isinstance(content, list):
                                        for item in content:
                                            if (
                                                isinstance(item, dict)
                                                and item.get("type") == "text"
                                            ):
                                                text = item.get("text", "").strip()

                                                # Skip tool results
                                                if text.startswith("tool_use_id"):
                                                    continue

                                                # Skip interruption messages
                                                if "[Request interrupted" in text:
                                                    continue

                                                # Skip Claude's session continuation messages
                                                if "session is being continued" in text.lower():
                                                    continue

                                                # Remove XML-like tags (command messages, etc)
                                                import re

                                                text = re.sub(r"<[^>]+>", "", text).strip()

                                                # Skip command outputs
                                                if "is running" in text and "…" in text:
                                                    continue

                                                # Handle image references - extract text after them
                                                if text.startswith("[Image #"):
                                                    parts = text.split("]", 1)
                                                    if len(parts) > 1:
                                                        text = parts[1].strip()

                                                # If we have real user text, use it
                                                if (
                                                    text and len(text) > MIN_PREVIEW_TEXT_LENGTH
                                                ):  # Lower threshold to catch "hello"
                                                    first_user_msg = text[
                                                        :PREVIEW_TEXT_TRUNCATE_LENGTH
                                                    ].replace("\n", " ")
                                                    break

                                    # Handle string content (less common but possible)
                                    elif isinstance(content, str):
                                        import re

                                        content = content.strip()

                                        # Remove XML-like tags
                                        content = re.sub(r"<[^>]+>", "", content).strip()

                                        # Skip command outputs
                                        if "is running" in content and "…" in content:
                                            continue

                                        # Skip Claude's session continuation messages
                                        if "session is being continued" in content.lower():
                                            continue

                                        # Skip tool results and interruptions
                                        if (
                                            not content.startswith("tool_use_id")
                                            and "[Request interrupted" not in content
                                        ):
                                            if (
                                                content and len(content) > MIN_PREVIEW_TEXT_LENGTH
                                            ):  # Lower threshold to catch short messages
                                                first_user_msg = content[
                                                    :PREVIEW_TEXT_TRUNCATE_LENGTH
                                                ].replace("\n", " ")
                        except json.JSONDecodeError:
                            continue

            return first_user_msg or "No preview available", msg_count
        except Exception as e:
            return f"Error: {str(e)[:PREVIEW_ERROR_TRUNCATE_LENGTH]}", 0

    def list_recent_sessions(self, limit: int = None) -> List[Path]:
        """List recent sessions with details."""
        sessions = self.find_sessions()

        if not sessions:
            print("❌ No Claude sessions found in ~/.claude/projects/")
            print("💡 Make sure you've used Claude Code and have conversations saved.")
            return []

        print(f"\n📚 Found {len(sessions)} Claude sessions:\n")
        print("=" * LIST_SEPARATOR_WIDTH)

        # Show all sessions if no limit specified
        sessions_to_show = sessions[:limit] if limit else sessions
        for i, session in enumerate(sessions_to_show, 1):
            # Clean up project name (remove hyphens, make readable)
            project = session.parent.name.replace("-", " ").strip()
            if project.startswith("Users"):
                project = (
                    "~/" + "/".join(project.split()[2:]) if len(project.split()) > 2 else "Home"
                )

            session_id = session.stem
            modified = datetime.fromtimestamp(session.stat().st_mtime)

            # Get file size
            size = session.stat().st_size
            size_kb = size / 1024

            # Get preview and message count
            preview, msg_count = self.get_conversation_preview(session)

            # Print formatted info
            print(f"\n{i}. 📁 {project}")
            print(f"   📄 Session: {session_id[:SESSION_ID_MAX_LENGTH]}...")
            print(f"   📅 Modified: {modified.strftime('%Y-%m-%d %H:%M')}")
            print(f"   💬 Messages: {msg_count}")
            print(f"   💾 Size: {size_kb:.1f} KB")
            print(f'   📝 Preview: "{preview}..."')

        print("\n" + "=" * LIST_SEPARATOR_WIDTH)
        return sessions[:limit]

    def extract_multiple(
        self,
        sessions: List[Path],
        indices: List[int],
        format: str = "markdown",
        detailed: bool = False,
    ) -> Tuple[int, int]:
        """Extract multiple sessions by index.

        Args:
            sessions: List of session paths
            indices: Indices to extract
            format: Output format ('markdown', 'json', 'html')
            detailed: If True, include tool use and system messages
        """
        success = 0
        total = len(indices)

        for idx in indices:
            if 0 <= idx < len(sessions):
                session_path = sessions[idx]
                # Extract project name from session path
                project_name = session_path.parent.name
                conversation = self.extract_conversation(session_path, detailed=detailed)
                if conversation:
                    output_path = self.save_conversation(
                        conversation, session_path.stem, format=format, project_name=project_name
                    )
                    success += 1
                    msg_count = len(conversation)
                    print(f"✅ {success}/{total}: {output_path.name} " f"({msg_count} messages)")
                else:
                    print(f"⏭️  Skipped session {idx + 1} (no conversation)")
            else:
                print(f"❌ Invalid session number: {idx + 1}")

        return success, total


def main():
    parser = argparse.ArgumentParser(
        description="Extract Claude Code conversations to clean markdown files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --list                    # List all available sessions
  %(prog)s --extract 1               # Extract the most recent session
  %(prog)s --extract 1,3,5           # Extract specific sessions
  %(prog)s --recent 5                # Extract 5 most recent sessions
  %(prog)s --all                     # Extract all sessions
  %(prog)s --output ~/my-logs        # Specify output directory
  %(prog)s --search "python error"   # Search conversations
  %(prog)s --search-regex "import.*" # Search with regex
  %(prog)s --format json --all       # Export all as JSON
  %(prog)s --format html --extract 1 # Export session 1 as HTML
  %(prog)s --format pdf --extract 1  # Export as PDF (requires: pip install .[pdf])
  %(prog)s --format docx --all       # Export as DOCX (requires: pip install .[docx])
  %(prog)s --detailed --extract 1    # Include tool use & system messages
        """,
    )
    parser.add_argument("--list", action="store_true", help="List recent sessions")
    parser.add_argument(
        "--extract",
        type=str,
        help="Extract specific session(s) by number (comma-separated)",
    )
    parser.add_argument("--all", "--logs", action="store_true", help="Extract all sessions")
    parser.add_argument("--recent", type=int, help="Extract N most recent sessions", default=0)
    parser.add_argument("--output", type=str, help="Output directory for markdown files")
    parser.add_argument(
        "--limit", type=int, help="Limit for --list command (default: show all)", default=None
    )
    parser.add_argument(
        "--interactive",
        "-i",
        "--start",
        "-s",
        action="store_true",
        help="Launch interactive UI for easy extraction",
    )
    parser.add_argument(
        "--export",
        type=str,
        help="Export mode: 'logs' for interactive UI",
    )

    # Search arguments
    parser.add_argument("--search", type=str, help="Search conversations for text (smart search)")
    parser.add_argument("--search-regex", type=str, help="Search conversations using regex pattern")
    parser.add_argument("--search-date-from", type=str, help="Filter search from date (YYYY-MM-DD)")
    parser.add_argument("--search-date-to", type=str, help="Filter search to date (YYYY-MM-DD)")
    parser.add_argument(
        "--search-speaker",
        choices=["human", "assistant", "both"],
        default="both",
        help="Filter search by speaker",
    )
    parser.add_argument("--case-sensitive", action="store_true", help="Make search case-sensitive")

    # Export format arguments
    parser.add_argument(
        "--format",
        choices=["markdown", "json", "html", "pdf", "docx"],
        default="markdown",
        help="Output format for exported conversations (default: markdown). PDF and DOCX require optional dependencies.",
    )
    parser.add_argument(
        "--detailed",
        action="store_true",
        help="Include tool use, MCP responses, and system messages in export",
    )

    args = parser.parse_args()

    # Handle interactive mode
    if args.interactive or (args.export and args.export.lower() == "logs"):
        from interactive_ui import main as interactive_main

        interactive_main()
        return

    # Initialize extractor with optional output directory
    extractor = ClaudeConversationExtractor(args.output)

    # Handle search mode
    if args.search or args.search_regex:
        from datetime import datetime

        from search_conversations import ConversationSearcher

        searcher = ConversationSearcher()

        # Determine search mode and query
        if args.search_regex:
            query = args.search_regex
            mode = "regex"
        else:
            query = args.search
            mode = "smart"

        # Parse date filters
        date_from = None
        date_to = None
        if args.search_date_from:
            try:
                date_from = datetime.strptime(args.search_date_from, "%Y-%m-%d")
            except ValueError:
                print(f"❌ Invalid date format: {args.search_date_from}")
                return

        if args.search_date_to:
            try:
                date_to = datetime.strptime(args.search_date_to, "%Y-%m-%d")
            except ValueError:
                print(f"❌ Invalid date format: {args.search_date_to}")
                return

        # Speaker filter
        speaker_filter = None if args.search_speaker == "both" else args.search_speaker

        # Perform search
        print(f"🔍 Searching for: {query}")
        results = searcher.search(
            query=query,
            mode=mode,
            date_from=date_from,
            date_to=date_to,
            speaker_filter=speaker_filter,
            case_sensitive=args.case_sensitive,
            max_results=SEARCH_MAX_RESULTS_DEFAULT,
        )

        if not results:
            print("❌ No matches found.")
            return

        print(f"\n✅ Found {len(results)} matches across conversations:")

        # Group and display results
        results_by_file = {}
        for result in results:
            if result.file_path not in results_by_file:
                results_by_file[result.file_path] = []
            results_by_file[result.file_path].append(result)

        # Store file paths for potential viewing
        file_paths_list = []
        for file_path, file_results in results_by_file.items():
            file_paths_list.append(file_path)
            print(
                f"\n{len(file_paths_list)}. 📄 {file_path.parent.name} ({len(file_results)} matches)"
            )
            # Show first match preview
            first = file_results[0]
            print(f"   {first.speaker}: {first.matched_content[:PREVIEW_TEXT_TRUNCATE_LENGTH]}...")

        # Offer to view conversations
        if file_paths_list:
            print("\n" + "=" * MAJOR_SEPARATOR_WIDTH)
            try:
                view_choice = input(
                    "\nView a conversation? Enter number (1-{}) or press Enter to skip: ".format(
                        len(file_paths_list)
                    )
                ).strip()

                if view_choice.isdigit():
                    view_num = int(view_choice)
                    if 1 <= view_num <= len(file_paths_list):
                        selected_path = file_paths_list[view_num - 1]
                        extractor.display_conversation(selected_path, detailed=args.detailed)

                        # Offer to extract after viewing
                        extract_choice = (
                            input("\n📤 Extract this conversation? (y/N): ").strip().lower()
                        )
                        if extract_choice == "y":
                            conversation = extractor.extract_conversation(
                                selected_path, detailed=args.detailed
                            )
                            if conversation:
                                session_id = selected_path.stem
                                project_name = selected_path.parent.name
                                if args.format == "json":
                                    output = extractor.save_as_json(
                                        conversation, session_id, project_name
                                    )
                                elif args.format == "html":
                                    output = extractor.save_as_html(
                                        conversation, session_id, project_name
                                    )
                                else:
                                    output = extractor.save_as_markdown(
                                        conversation, session_id, project_name
                                    )
                                print(f"✅ Saved: {output.name}")
            except (EOFError, KeyboardInterrupt):
                print("\n👋 Cancelled")

        return

    # Default action is to list sessions
    if args.list or (
        not args.extract
        and not args.all
        and not args.recent
        and not args.search
        and not args.search_regex
    ):
        sessions = extractor.list_recent_sessions(args.limit)

        if sessions and not args.list:
            print("\nTo extract conversations:")
            print("  claude-extract --extract <number>      # Extract specific session")
            print("  claude-extract --recent 5              # Extract 5 most recent")
            print("  claude-extract --all                   # Extract all sessions")

    elif args.extract:
        sessions = extractor.find_sessions()

        # Parse comma-separated indices
        indices = []
        for num in args.extract.split(","):
            try:
                idx = int(num.strip()) - 1  # Convert to 0-based index
                indices.append(idx)
            except ValueError:
                print(f"❌ Invalid session number: {num}")
                continue

        if indices:
            print(f"\n📤 Extracting {len(indices)} session(s) as {args.format.upper()}...")
            if args.detailed:
                print("📋 Including detailed tool use and system messages")
            success, total = extractor.extract_multiple(
                sessions, indices, format=args.format, detailed=args.detailed
            )
            print(f"\n✅ Successfully extracted {success}/{total} sessions")

    elif args.recent:
        sessions = extractor.find_sessions()
        limit = min(args.recent, len(sessions))
        print(f"\n📤 Extracting {limit} most recent sessions as {args.format.upper()}...")
        if args.detailed:
            print("📋 Including detailed tool use and system messages")

        indices = list(range(limit))
        success, total = extractor.extract_multiple(
            sessions, indices, format=args.format, detailed=args.detailed
        )
        print(f"\n✅ Successfully extracted {success}/{total} sessions")

    elif args.all:
        sessions = extractor.find_sessions()
        print(f"\n📤 Extracting all {len(sessions)} sessions as {args.format.upper()}...")
        if args.detailed:
            print("📋 Including detailed tool use and system messages")

        indices = list(range(len(sessions)))
        success, total = extractor.extract_multiple(
            sessions, indices, format=args.format, detailed=args.detailed
        )
        print(f"\n✅ Successfully extracted {success}/{total} sessions")


def launch_interactive():
    """Launch the interactive UI directly, or handle search if specified."""
    import sys

    # If no arguments provided, launch interactive UI
    if len(sys.argv) == 1:
        try:
            from .interactive_ui import main as interactive_main
        except ImportError:
            from interactive_ui import main as interactive_main
        interactive_main()
    else:
        # If other arguments are provided, run the normal CLI
        main()


if __name__ == "__main__":
    main()
